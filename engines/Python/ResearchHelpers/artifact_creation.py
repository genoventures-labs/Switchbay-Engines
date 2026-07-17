"""
Artifact Creation Engine

Artifact Creation Engine for Switchbay.
This engine is designed to facilitate the creation of rich, structured research artifacts and reports
across multiple file formats, serving both as output for automated agents/models and assisting human users.

Supported Output Types:
- Markdown (.md): Full support for metadata/frontmatter, multi-section structure.
- PDF (.pdf): Styled research PDFs with title, author, and multi-level sections (via ReportLab).
- DOCX (.docx): Structured Word documents with sectioning, metadata, and optional title/author (python-docx).
- HTML (.html): Self-contained HTML reports/pages with meta tags, title, and structured content.
- Plain Text (.txt): Simple or structured human-readable text, with optional headers/metadata and sectioning.

Features:
- Add descriptive metadata (title, author, custom fields).
- Multi-section support for outlining, appendices, references, methodologies, etc.
- Consistent CLI and Python API for report generation.
- Useful for automated research agents and reproducible scientific documentation.

Example CLI Usage:
  python engines/Python/ResearchHelpers/artifact_creation.py create_markdown --content "..." --filename "out.md" [--title "..."] [--metadata "key: val"] [--sections "[{'title':'Intro','body':'...'}]"]
  python engines/Python/ResearchHelpers/artifact_creation.py create_pdf --content "..." --filename "out.pdf" [--title "..."] [--author "..."] [--sections ...]
  python engines/Python/ResearchHelpers/artifact_creation.py create_docx --content "..." --filename "out.docx" [--title "..."] [--author "..."] [--sections ...]
  python engines/Python/ResearchHelpers/artifact_creation.py create_html --content "..." --filename "out.html" [--title "..."] [--metadata ...] [--sections ...]
  python engines/Python/ResearchHelpers/artifact_creation.py create_txt --content "..." --filename "out.txt" [--metadata ...] [--sections ...]
"""

from __future__ import annotations

import argparse
import os
import sys
import time
import urllib.parse
from typing import Any, Dict, List, Optional
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen



def create_markdown(
    content: str,
    filename: str = "artifact.md",
    metadata: Optional[Dict[str, Any]] = None,
    sections: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """
    Create a structured markdown research artifact or report.

    Args:
        content (str): The main markdown body content.
        filename (str): Name of the output file.
        metadata (dict): Optional frontmatter metadata (dict will be dumped as YAML frontmatter).
        sections (list): Optional list of sections, each as {'title':..., 'body':...}.

    Returns:
        dict: Result status and filename.
    """
    out = []
    if metadata:
        try:
            import yaml  # type: ignore

            out.append("---")
            out.append(yaml.safe_dump(metadata, sort_keys=False).strip())
            out.append("---\n")
        except ImportError:
            # Fallback without PyYAML
            out.append("---")
            for k, v in metadata.items():
                out.append(f"{k}: {v}")
            out.append("---\n")
    if sections:
        for section in sections:
            if section.get("title"):
                out.append(f"# {section['title']}\n")
            if section.get("body"):
                out.append(section["body"] + "\n")
    else:
        out.append(content)
    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    return {"success": True, "filename": filename}

def create_pdf(
    content: str,
    filename: str = "artifact.pdf",
    title: Optional[str] = None,
    author: Optional[str] = None,
    sections: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """
    Create a PDF research artifact, with optional metadata and sectioning.

    Args:
        content (str): Main content for the PDF (if sections is not supplied).
        filename (str): Output file name.
        title (str): Optional title for the PDF.
        author (str): Optional author(s).
        sections (list): Optional list of sections as {'title':..., 'body':...}

    Returns:
        dict: Status and filename.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        y = height - 50

        if title:
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, y, title)
            y -= 28
        if author:
            c.setFont("Helvetica-Oblique", 10)
            c.drawString(72, y, f"Author: {author}")
            y -= 18
        y -= 18

        c.setFont("Helvetica", 11)
        def draw_wrap(text, x, y, c, max_width=width-100):
            from reportlab.lib.utils import simpleSplit
            parts = simpleSplit(text, "Helvetica", 11, max_width)
            for part in parts:
                c.drawString(x, y, part)
                y -= 16
            return y

        if sections:
            for section in sections:
                if section.get("title"):
                    c.setFont("Helvetica-Bold", 13)
                    y -= 8
                    c.drawString(72, y, section["title"])
                    y -= 20
                    c.setFont("Helvetica", 11)
                if section.get("body"):
                    y = draw_wrap(section["body"], 80, y, c)
                    y -= 10
                if y < 60:
                    c.showPage()
                    y = height - 50
        else:
            y = draw_wrap(content, 72, y, c)

        c.save()
        return {"success": True, "filename": filename}
    except ImportError:
        return {"success": False, "error": "Missing reportlab dependency. Run 'pip install reportlab'."}

def create_docx(
    content: str,
    filename: str = "artifact.docx",
    title: Optional[str] = None,
    author: Optional[str] = None,
    sections: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """
    Create a DOCX research artifact with optional sections, title, and author.

    Args:
        content (str): Main content if no sections supplied.
        filename (str): Output file name.
        title (str): Optional report/document title.
        author (str): Optional author metadata.
        sections (list): List of {'title':..., 'body':...}

    Returns:
        dict: Status and filename.
    """
    try:
        from docx import Document

        doc = Document()
        if title:
            doc.add_heading(title, 0)
        if author:
            doc.add_paragraph(f"Author: {author}")

        if sections:
            for section in sections:
                if section.get("title"):
                    doc.add_heading(section["title"], level=1)
                if section.get("body"):
                    doc.add_paragraph(section["body"])
        else:
            doc.add_paragraph(content)

        doc.save(filename)
        return {"success": True, "filename": filename}
    except ImportError:
        return {"success": False, "error": "Missing python-docx dependency. Run 'pip install python-docx'."}

def create_html(
    content: str,
    filename: str = "artifact.html",
    title: Optional[str] = None,
    metadata: Optional[Dict[str, str]] = None,
    sections: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """
    Create an HTML research artifact with basic structure, optional metadata and sectioning.

    Args:
        content (str): Main HTML content if no sections supplied.
        filename (str): Output HTML filename.
        title (str): Optional HTML title element.
        metadata (dict): Optional metadata headers (added as meta tags).
        sections (list): Optional [{'title':..., 'body':...}]
    Returns:
        dict: Status and filename.
    """
    html_lines = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>"
    ]
    html_lines.append(f"<meta charset='utf-8'>")
    if title:
        html_lines.append(f"<title>{title}</title>")
    if metadata:
        for k, v in metadata.items():
            html_lines.append(f"<meta name='{k}' content='{v}'>")
    html_lines.append("</head>")
    html_lines.append("<body>")
    if title:
        html_lines.append(f"<h1>{title}</h1>")
    if sections:
        for section in sections:
            if section.get("title"):
                html_lines.append(f"<h2>{section['title']}</h2>")
            if section.get("body"):
                html_lines.append(f"<p>{section['body']}</p>")
    else:
        html_lines.append(content)
    html_lines.append("</body></html>")

    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(html_lines))
    return {"success": True, "filename": filename}

def create_txt(
    content: str,
    filename: str = "artifact.txt",
    sections: Optional[List[Dict[str, str]]] = None,
    metadata: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    """
    Create a plain TXT research artifact, with optional header/sections/metadata as text.

    Args:
        content (str): Main TEXT content.
        filename (str): Output file name.
        sections (list): Optional list of {'title':..., 'body':...}.
        metadata (dict): Optional: Add text metadata header.

    Returns:
        dict: Status and filename.
    """
    lines = []
    if metadata:
        lines.append("# Metadata")
        for k, v in metadata.items():
            lines.append(f"{k}: {v}")
        lines.append("")
    if sections:
        for section in sections:
            if section.get("title"):
                lines.append(f"{section['title']}")
            if section.get("body"):
                lines.append(section["body"])
            lines.append("")
    else:
        lines.append(content)
    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return {"success": True, "filename": filename}

def _cli():
    parser = argparse.ArgumentParser(description='Artifact Creation Engine: supports Markdown, PDF, DOCX, HTML, TXT with metadata and structure. Use --help for details.')

    parser.add_argument('mode', type=str, help='Artifact type: create_markdown, create_pdf, create_docx, create_html, create_txt')
    parser.add_argument('--content', type=str, required=False, help='Main content to include in the artifact.')
    parser.add_argument('--filename', type=str, required=True, help='Filename to create the artifact as')
    parser.add_argument('--title', type=str, required=False, help='Title for the artifact/report (where supported)')
    parser.add_argument('--author', type=str, required=False, help='Author of the artifact (DOCX, PDF)')
    parser.add_argument('--metadata', type=str, required=False, help='YAML string of metadata for markdown/html/txt headers.')
    parser.add_argument('--sections', type=str, required=False, help='JSON list of sections, e.g. [{"title": "Intro", "body": "..."}]')

    args = parser.parse_args()

    import json

    try:
        import yaml  # type: ignore
    except ImportError:
        yaml = None  # type: ignore

    metadata = None
    if args.metadata:
        if yaml is not None:
            try:
                metadata = yaml.safe_load(args.metadata)
            except Exception:
                metadata = None
        if metadata is None:
            try:
                metadata = json.loads(args.metadata)
            except Exception:
                metadata = None

    sections = None
    if args.sections:
        try:
            sections = json.loads(args.sections)
        except Exception:
            print(json.dumps({"success": False, "error": "Failed to parse --sections as JSON."}))
            sys.exit(1)

    kwargs = dict(
        content=args.content or "",
        filename=args.filename,
        title=getattr(args, "title", None),
        author=getattr(args, "author", None),
        metadata=metadata,
        sections=sections,
    )

    # Remove keys not accepted by each function
    MODE_MAP = {
        "create_markdown": create_markdown,
        "create_pdf": create_pdf,
        "create_docx": create_docx,
        "create_html": create_html,
        "create_txt": create_txt,
    }
    fn = MODE_MAP.get(args.mode)
    if fn is None:
        print(f"Unknown mode {args.mode}.")
        sys.exit(1)

    # Remove arguments not accepted by target function
    import inspect
    sig = inspect.signature(fn)
    filtered = {k: v for k, v in kwargs.items() if k in sig.parameters}
    result = fn(**filtered)
    print(json.dumps(result, indent=2, ensure_ascii=False))
    if not result.get("success", True):
        sys.exit(1)


if __name__ == "__main__":
    _cli()
